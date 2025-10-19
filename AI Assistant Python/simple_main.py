from fastapi import FastAPI, Request, Form, Depends, UploadFile, File
from fastapi.templating import Jinja2Templates
from fastapi.responses import HTMLResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
import logging
from datetime import datetime
import uuid
from dotenv import load_dotenv
import pandas as pd
import io
from typing import List
import base64
import os

# Load environment variables from .env file
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Create FastAPI app
app = FastAPI(title="AI Duty Officer Assistant", version="1.0.0")

# Resolve paths relative to this file
script_dir = os.path.dirname(os.path.abspath(__file__))

# Setup static files with absolute path to avoid CWD issues
static_dir = os.path.join(script_dir, "static")
app.mount("/static", StaticFiles(directory=static_dir), name="static")

# Setup templates with correct path
templates_dir = os.path.join(script_dir, "app", "templates")
templates = Jinja2Templates(directory=templates_dir)

# Mock data classes for now
class MockIncident:
    def __init__(self, description, source="Manual"):
        self.id = str(uuid.uuid4())
        self.description = description
        self.source = source
        self.reported_at = datetime.now()
        self.status = "New"

# Import the real services
from app.services.openai_service import OpenAIService
from app.services.knowledge_base_service import KnowledgeBaseService
from app.services.training_data_service import TrainingDataService
from app.services.incident_analyzer import IncidentAnalyzer
from app.services.error_matcher_service import ErrorTypeMatcher
from app.services.document_parser_service import DocumentParserService
from app.models.database import Base
from app.database import get_db, engine
from sqlalchemy.orm import Session

# Create database tables
Base.metadata.create_all(bind=engine)

# Initialize the OpenAI service
openai_service = OpenAIService()

# Initialize the Document Parser service
document_parser = DocumentParserService(openai_service)

async def analyze_image_with_ai(image_content: bytes, content_type: str) -> str:
    """Analyze image using Azure OpenAI Vision API"""
    try:
        # Convert image to base64
        import base64
        encoded_image = base64.b64encode(image_content).decode('utf-8')
        
        # Use Azure OpenAI Vision to analyze the image
        vision_analysis = await openai_service.analyze_image_async(encoded_image, "Maritime incident documentation")
        
        return f"Visual Analysis: {vision_analysis} "
        
    except Exception as ex:
        logger.error(f"Error analyzing image: {ex}")
        return "[Image analysis failed] "

async def validate_incident_input(description: str) -> dict:
    """
    Use AI to validate if the input is a legitimate incident description
    Returns validation result with reasoning
    """
    try:
        # Basic checks first
        description = description.strip()
        
        # Empty or too short
        if len(description) < 5:
            return {
                "valid": False,
                "reason": "Description too short - please provide more details about the incident"
            }
        
        # Too long (potential spam/copy-paste)
        if len(description) > 5000:
            return {
                "valid": False, 
                "reason": "Description too long - please provide a concise incident summary"
            }
        
        # Use AI to validate content quality
        validation_prompt = f"""You are validating incident reports for a maritime operations system. Determine if this input is a legitimate incident description.

Input: {description[:800]}

VALID incident descriptions include:
- Technical problems (system errors, equipment failures)
- Operational issues (delays, process problems) 
- Safety concerns or incidents
- Infrastructure problems
- Service disruptions
- Detailed problem reports with context

INVALID inputs include:
- Random text or gibberish ("asdf", "test", "hello world")
- Single words or very short phrases without context
- Nonsensical combinations of words
- Personal messages not related to operations
- Jokes, memes, or casual conversation
- Testing inputs or placeholder text
- Spam or repeated characters

Respond with only: VALID or INVALID"""

        response = await openai_service.get_completion(
            messages=[{"role": "user", "content": validation_prompt}],
            max_tokens=10,
            temperature=0.05
        )
        
        validation = response.strip().upper()
        
        if validation == "INVALID":
            return {
                "valid": False,
                "reason": "Input appears to be random text or not a legitimate incident report. Please provide a clear description of an operational issue, technical problem, or safety concern."
            }
        
        return {"valid": True, "reason": "Input validated successfully"}
        
    except Exception as e:
        logger.error(f"Error validating input: {e}")
        # If validation fails, be conservative and allow through
        return {"valid": True, "reason": "Validation service unavailable - input accepted"}

class MockResolutionStep:
    def __init__(self, order, description, step_type="Analysis"):
        self.order = order
        self.description = description
        self.type = step_type
        self.query = ""

class MockResolutionPlan:
    def __init__(self, incident_type):
        self.summary = f"Analysis completed for {incident_type}"
        self.steps = [
            MockResolutionStep(1, "Initial assessment completed using AI analysis", "Analysis"),
            MockResolutionStep(2, "Investigate root cause based on analysis findings", "Investigation"),
            MockResolutionStep(3, "Implement resolution based on findings", "Resolution")
        ]
        self.diagnostic_queries = []
        self.resolution_queries = []

# Routes
@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    """Home page"""
    return templates.TemplateResponse("index.html", {"request": request})

@app.get("/analyze", response_class=HTMLResponse)
async def analyze_get(request: Request):
    """Analyze page - GET"""
    test_cases = [
        {
            "description": "Customer on PORTNET is seeing 2 identical containers information for CMAU0000020",
            "source": "Email",
            "priority": "Medium", 
            "title": "Container Duplication Issue",
            "icon": "fas fa-box",
            "category": "Container Management"
        },
        {
            "description": "VESSEL_ERR_4 when creating vessel advice for MV Lion City 07",
            "source": "Email",
            "priority": "High",
            "title": "Vessel Operations Error", 
            "icon": "fas fa-ship",
            "category": "Vessel Operations"
        },
        {
            "description": "EDI message REF-IFT-0007 stuck in ERROR status, ack_at is NULL",
            "source": "SMS",
            "priority": "High",
            "title": "EDI Processing Failure",
            "icon": "fas fa-exchange-alt",
            "category": "Data Integration"
        }
    ]
    
    return templates.TemplateResponse("analyze.html", {
        "request": request, 
        "test_cases": test_cases
    })

@app.post("/analyze")
async def analyze_post(
    request: Request,
    incident_description: str = Form(...),
    incident_source: str = Form("Manual"),
    incident_images: List[UploadFile] = File(default=[]),
    log_files: List[UploadFile] = File(default=[]),
    db: Session = Depends(get_db)
):
    """Extract information from incident description for user review"""
    try:
        # Validate incident input first
        validation_result = await validate_incident_input(incident_description)

        
        if not validation_result["valid"]:
            logger.warning(f"Invalid incident input rejected: {incident_description[:50]}...")
            error_message = validation_result["reason"]
            return templates.TemplateResponse("analyze.html", {
                "request": request,
                "error": error_message,
                "description": incident_description,
                "test_cases": [
                    {
                        "title": "Container System Error",
                        "description": "PORTNET container CMAU123456 showing duplicate entries causing discharge delays"
                    },
                    {
                        "title": "EDI Processing Failure", 
                        "description": "EDI message IFT-001 stuck in ERROR status, ack_at field NULL for 2 hours"
                    },
                    {
                        "title": "Vessel Berth Issue",
                        "description": "MV Pacific Star unable to berth at Terminal 3 due to equipment malfunction"
                    }
                ]
            })
        
        logger.info(f"Input validation passed: {validation_result['reason']}")
        
        # Store uploaded files temporarily in session or save them
        has_images = bool(incident_images and incident_images[0].filename)
        has_logs = bool(log_files and log_files[0].filename)
        
        # Process uploaded files if any
        if has_images:
            logger.info(f"Processing {len(incident_images)} uploaded images")
            
            # Create uploads directory if it doesn't exist
            import os
            uploads_dir = os.path.join(os.path.dirname(__file__), "static", "uploads")
            os.makedirs(uploads_dir, exist_ok=True)
            
            for image in incident_images:
                if image.filename and image.content_type.startswith('image/'):
                    # Save image temporarily
                    import uuid
                    file_extension = os.path.splitext(image.filename)[1]
                    unique_filename = f"temp_{uuid.uuid4()}{file_extension}"
                    file_path = os.path.join(uploads_dir, unique_filename)
                    
                    content = await image.read()
                    with open(file_path, "wb") as f:
                        f.write(content)
        
        # Extract structured information using AI
        extracted_info = await openai_service.extract_incident_information(incident_description)
        logger.info(f"Extracted incident information: {list(extracted_info.keys())}")
        
        # Render extraction review template
        return templates.TemplateResponse("extraction_review.html", {
            "request": request,
            "extracted_info": extracted_info,
            "original_description": incident_description,
            "incident_source": incident_source,
            "has_images": has_images,
            "has_logs": has_logs
        })
        
    except Exception as ex:
        logger.error(f"Error in information extraction: {ex}")
        return templates.TemplateResponse("analyze.html", {
            "request": request,
            "error": f"Error extracting information: {str(ex)}",
            "description": incident_description,
            "test_cases": []
        })

@app.post("/analyze-confirmed")
async def analyze_confirmed(
    request: Request,
    original_description: str = Form(...),
    incident_source: str = Form("Manual"),
    has_images: str = Form(default=""),
    has_logs: str = Form(default=""),
    # Extracted and user-reviewed fields
    incident_date: str = Form(...),
    location: str = Form(...),
    vessel_name: str = Form(...),
    vessel_type: str = Form(...),
    vessel_flag: str = Form(...),
    incident_type: str = Form(...),
    severity_level: str = Form(...),
    weather_conditions: str = Form(...),
    personnel_involved: str = Form(...),
    injuries_fatalities: str = Form(...),
    equipment_involved: str = Form(...),
    cargo_details: str = Form(...),
    immediate_actions: str = Form(...),
    estimated_damage: str = Form(...),
    authorities_notified: str = Form(...),
    environmental_impact: str = Form(...),
    db: Session = Depends(get_db)
):
    """Perform full incident analysis with confirmed information"""
    try:
        # Create enhanced description from structured information
        structured_info = {
            "original_description": original_description,
            "incident_date": incident_date,
            "location": location,
            "vessel_name": vessel_name,
            "vessel_type": vessel_type,
            "vessel_flag": vessel_flag,
            "incident_type": incident_type,
            "severity_level": severity_level,
            "weather_conditions": weather_conditions,
            "personnel_involved": personnel_involved,
            "injuries_fatalities": injuries_fatalities,
            "equipment_involved": equipment_involved,
            "cargo_details": cargo_details,
            "immediate_actions": immediate_actions,
            "estimated_damage": estimated_damage,
            "authorities_notified": authorities_notified,
            "environmental_impact": environmental_impact
        }
        
        # Create enhanced description that includes structured information
        enhanced_description = f"""
INCIDENT DESCRIPTION:
{original_description}

STRUCTURED INFORMATION SUMMARY:
Date/Time: {incident_date}
Location: {location}
Vessel: {vessel_name} ({vessel_type}, Flag: {vessel_flag})
Incident Type: {incident_type}
Severity: {severity_level}
Weather: {weather_conditions}
Personnel: {personnel_involved}
Injuries/Fatalities: {injuries_fatalities}
Equipment: {equipment_involved}
Cargo: {cargo_details}
Immediate Actions: {immediate_actions}
Damage: {estimated_damage}
Authorities: {authorities_notified}
Environmental Impact: {environmental_impact}
"""
        
        # Process any uploaded files if indicated
        image_analysis = ""
        uploaded_images = []
        
        if has_images == "true":
            # Find and process temporary images
            import os
            import glob
            uploads_dir = os.path.join(os.path.dirname(__file__), "static", "uploads")
            temp_files = glob.glob(os.path.join(uploads_dir, "temp_*"))
            
            for file_path in temp_files:
                if os.path.exists(file_path):
                    filename = os.path.basename(file_path)
                    # Remove temp_ prefix for display
                    display_name = filename.replace("temp_", "")
                    uploaded_images.append({
                        "filename": filename,
                        "original_name": display_name,
                        "path": f"/static/uploads/{filename}",
                        "size": os.path.getsize(file_path)
                    })
                    
                    # Analyze image with AI vision
                    try:
                        with open(file_path, "rb") as f:
                            content = f.read()
                        image_analysis += await analyze_image_with_ai(content, "image/jpeg")
                    except Exception as e:
                        logger.error(f"Error analyzing image {filename}: {e}")
        
        # Combine enhanced description with image analysis
        combined_description = enhanced_description
        if image_analysis:
            combined_description += f"\n\nIMAGE ANALYSIS:\n{image_analysis}"
        
        # Create mock incident
        incident = MockIncident(combined_description, incident_source)
        
        # Use the full IncidentAnalyzer that includes knowledge base integration
        incident_analyzer = IncidentAnalyzer(db)
        analysis = await incident_analyzer.analyze_incident_async(combined_description)
        
        # Generate AI-powered resolution plan
        resolution_data = await openai_service.generate_resolution_plan_async(combined_description, analysis)
        
        # Convert to expected format
        class AIResolutionStep:
            def __init__(self, step_data):
                self.order = step_data.get("order", 1)
                self.description = step_data.get("description", "")
                self.type = step_data.get("type", "Analysis")
                self.query = step_data.get("query", "")
        
        class AIResolutionPlan:
            def __init__(self, resolution_data):
                self.summary = resolution_data.get("summary", "")
                self.steps = [AIResolutionStep(step) for step in resolution_data.get("steps", [])]
                self.diagnostic_queries = []
                self.resolution_queries = []
        
        resolution_plan = AIResolutionPlan(resolution_data)
        
        # Create mock view model with enhanced information
        class MockViewModel:
            def __init__(self, incident, analysis, resolution_plan, uploaded_images=None, structured_info=None):
                self.incident = incident
                self.analysis = analysis  
                self.resolution_plan = resolution_plan
                self.uploaded_images = uploaded_images or []
                self.structured_info = structured_info or {}
        
        view_model = MockViewModel(incident, analysis, resolution_plan, uploaded_images, structured_info)
        
        # Clean up temporary files
        if has_images == "true":
            import glob
            uploads_dir = os.path.join(os.path.dirname(__file__), "static", "uploads")
            temp_files = glob.glob(os.path.join(uploads_dir, "temp_*"))
            for file_path in temp_files:
                try:
                    os.remove(file_path)
                except Exception as e:
                    logger.warning(f"Could not remove temp file {file_path}: {e}")
        
        return templates.TemplateResponse("results.html", {
            "request": request,
            "result": view_model,
            "uploaded_images": uploaded_images
        })
        
    except Exception as ex:
        logger.error(f"Error analyzing incident: {ex}")
        return RedirectResponse(url="/analyze?error=Analysis failed", status_code=302)        

@app.get("/upload-knowledge")
async def upload_knowledge_get(request: Request):
    """Knowledge upload page"""
    return templates.TemplateResponse("upload_knowledge.html", {"request": request})

@app.get("/knowledge")
async def view_knowledge(request: Request, db: Session = Depends(get_db)):
    """View knowledge base entries"""
    try:
        knowledge_service = KnowledgeBaseService(db)
        entries = knowledge_service.get_all_knowledge(skip=0, limit=100)
        
        return templates.TemplateResponse("knowledge_list.html", {
            "request": request,
            "entries": entries
        })
    except Exception as ex:
        logger.error(f"Error retrieving knowledge: {ex}")
        return templates.TemplateResponse("knowledge_list.html", {
            "request": request,
            "entries": [],
            "error": f"Error loading knowledge base: {str(ex)}"
        })

@app.get("/training")
async def view_training(request: Request, db: Session = Depends(get_db)):
    """View training data entries"""
    try:
        from app.models.database import TrainingData
        training_data = db.query(TrainingData).order_by(TrainingData.created_at.desc()).all()
        
        return templates.TemplateResponse("training.html", {
            "request": request,
            "training_data": training_data
        })
    except Exception as ex:
        logger.error(f"Error retrieving training data: {ex}")
        return templates.TemplateResponse("training.html", {
            "request": request,
            "training_data": [],
            "error": f"Error loading training data: {str(ex)}"
        })

@app.get("/database-status")
async def database_status(request: Request, db: Session = Depends(get_db)):
    """View database status and contents"""
    try:
        from app.models.database import KnowledgeBase, TrainingData
        
        # Count entries
        kb_count = db.query(KnowledgeBase).count()
        td_count = db.query(TrainingData).count()
        
        # Get recent knowledge entries
        recent_knowledge = db.query(KnowledgeBase).order_by(KnowledgeBase.created_at.desc()).limit(10).all()
        
        # Get recent training data
        recent_training = db.query(TrainingData).order_by(TrainingData.created_at.desc()).limit(5).all()
        
        return templates.TemplateResponse("database_status.html", {
            "request": request,
            "kb_count": kb_count,
            "td_count": td_count,
            "recent_knowledge": recent_knowledge,
            "recent_training": recent_training
        })
    except Exception as ex:
        logger.error(f"Error retrieving database status: {ex}")
        return {"error": str(ex)}

@app.get("/sql-export")
async def sql_export(request: Request):
    """Export database as SQL"""
    try:
        import sqlite3
        
        # Connect to database
        conn = sqlite3.connect('duty_officer_assistant.db')
        cursor = conn.cursor()
        
        sql_content = []
        sql_content.append("-- =====================================================")
        sql_content.append("-- DUTY OFFICER ASSISTANT DATABASE EXPORT")
        sql_content.append(f"-- Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        sql_content.append("-- =====================================================\n")
        
        # Get table schemas and data
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = cursor.fetchall()
        
        for table_name in tables:
            table = table_name[0]
            sql_content.append(f"\n-- ===== TABLE: {table.upper()} =====")
            
            # Get table schema
            cursor.execute(f"SELECT sql FROM sqlite_master WHERE name='{table}'")
            schema = cursor.fetchone()
            if schema:
                sql_content.append(f"-- Schema:")
                sql_content.append(schema[0] + ";")
                sql_content.append("")
            
            # Get table data count
            cursor.execute(f"SELECT COUNT(*) FROM {table}")
            count = cursor.fetchone()[0]
            sql_content.append(f"-- Records: {count}")
            
            if count > 0:
                # Get column names
                cursor.execute(f"PRAGMA table_info({table})")
                columns = [col[1] for col in cursor.fetchall()]
                
                # Get all data
                cursor.execute(f"SELECT * FROM {table}")
                rows = cursor.fetchall()
                
                sql_content.append(f"\n-- Data for {table}:")
                for i, row in enumerate(rows):
                    insert_values = []
                    for value in row:
                        if value is None:
                            insert_values.append("NULL")
                        elif isinstance(value, str):
                            escaped_value = value.replace("'", "''")
                            insert_values.append(f"'{escaped_value}'")
                        else:
                            insert_values.append(str(value))
                    
                    column_list = "(" + ", ".join(columns) + ")"
                    values_list = "(" + ", ".join(insert_values) + ")"
                    sql_content.append(f"INSERT INTO {table} {column_list}")
                    sql_content.append(f"VALUES {values_list};")
                    sql_content.append("")
            
            sql_content.append(f"-- End of {table.upper()}")
            sql_content.append("-" * 60)
        
        conn.close()
        
        # Return as plain text response
        from fastapi.responses import PlainTextResponse
        return PlainTextResponse("\n".join(sql_content), media_type="text/plain")
        
    except Exception as ex:
        logger.error(f"Error exporting SQL: {ex}")
        return PlainTextResponse(f"Error exporting database: {str(ex)}", media_type="text/plain")

@app.post("/upload-knowledge")
async def upload_knowledge_post(
    request: Request, 
    title: str = Form(...), 
    category: str = Form(""), 
    content: str = Form(...),
    db: Session = Depends(get_db)
):
    """Handle knowledge upload"""
    try:
        # Use the real knowledge base service
        knowledge_service = KnowledgeBaseService(db)
        result = knowledge_service.import_from_word_content(
            content=content,
            title=title,
            category=category if category else "General",
            source="Web Upload"
        )
        
        logger.info(f"Knowledge uploaded successfully: {title} (ID: {result.id})")
        
        # Return success response
        return templates.TemplateResponse("upload_knowledge.html", {
            "request": request,
            "success": True,
            "message": f"Knowledge document '{title}' uploaded successfully! (ID: {result.id})"
        })
        
    except Exception as ex:
        logger.error(f"Error uploading knowledge: {ex}")
        return templates.TemplateResponse("upload_knowledge.html", {
            "request": request,
            "error": True,
            "message": f"Error uploading document: {str(ex)}"
        })

@app.get("/upload-training", response_class=HTMLResponse)
async def upload_training(request: Request):
    """Upload training data page"""
    return templates.TemplateResponse("upload_training.html", {"request": request})

@app.post("/upload-training-data")
async def upload_training_data(
    request: Request,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """Handle training data upload"""
    try:
        if not file.filename.endswith(('.xlsx', '.xls')):
            return templates.TemplateResponse("upload_training.html", {
                "request": request,
                "error": True,
                "message": "Please upload an Excel file (.xlsx or .xls)"
            })
        
        # Read Excel file
        content = await file.read()
        df = pd.read_excel(io.BytesIO(content))
        
        if df.empty:
            return templates.TemplateResponse("upload_training.html", {
                "request": request,
                "error": True,
                "message": "Excel file is empty"
            })
        
        # Intelligent column detection
        training_service = TrainingDataService(db)
        
        # Detect columns based on content patterns
        incident_col = None
        resolution_col = None
        
        # Look for incident-related columns
        for col in df.columns:
            col_lower = str(col).lower()
            sample_data = df[col].dropna().astype(str).str.lower()
            
            # Check if this looks like an incident column
            if any(keyword in col_lower for keyword in ['incident', 'problem', 'issue', 'description', 'summary', 'title']):
                incident_col = col
            # Check if this looks like a resolution column  
            elif any(keyword in col_lower for keyword in ['resolution', 'solution', 'fix', 'action', 'steps', 'procedure']):
                resolution_col = col
            # Content-based detection
            elif not incident_col and sample_data.str.contains('error|failed|down|issue|problem', na=False).any():
                incident_col = col
            elif not resolution_col and sample_data.str.contains('restart|check|verify|contact|replace', na=False).any():
                resolution_col = col
        
        # If no specific columns found, try first two text columns
        if not incident_col or not resolution_col:
            text_cols = [col for col in df.columns if df[col].dtype == 'object']
            if len(text_cols) >= 2:
                if not incident_col:
                    incident_col = text_cols[0]
                if not resolution_col:
                    resolution_col = text_cols[1]
            elif len(text_cols) == 1:
                incident_col = text_cols[0]
                resolution_col = text_cols[0]  # Use same column for both
        
        if not incident_col:
            return templates.TemplateResponse("upload_training.html", {
                "request": request,
                "error": True,
                "message": f"Could not identify incident column. Available columns: {list(df.columns)}"
            })
        
        # Process the data
        success_count = 0
        error_count = 0
        errors = []
        
        for index, row in df.iterrows():
            try:
                incident_text = str(row[incident_col]).strip()
                resolution_text = str(row[resolution_col]).strip() if resolution_col else ""
                
                if incident_text and incident_text.lower() not in ['nan', 'none', '']:
                    result = training_service.add_training_example(
                        incident_description=incident_text,
                        resolution_steps=resolution_text,
                        source=f"Excel Upload: {file.filename}",
                        category="Imported"
                    )
                    success_count += 1
                else:
                    error_count += 1
                    errors.append(f"Row {index + 1}: Empty incident description")
                    
            except Exception as ex:
                error_count += 1
                errors.append(f"Row {index + 1}: {str(ex)}")
        
        # Prepare result message
        message = f"Successfully imported {success_count} training examples"
        if incident_col:
            message += f" (Incident column: '{incident_col}'"
        if resolution_col and resolution_col != incident_col:
            message += f", Resolution column: '{resolution_col}'"
        if incident_col:
            message += ")"
        
        if error_count > 0:
            message += f". {error_count} errors occurred."
        
        return templates.TemplateResponse("upload_training.html", {
            "request": request,
            "success": True,
            "message": message,
            "details": {
                "success_count": success_count,
                "error_count": error_count,
                "errors": errors[:10],  # Show first 10 errors
                "incident_column": incident_col,
                "resolution_column": resolution_col,
                "total_rows": len(df)
            }
        })
        
    except Exception as ex:
        logger.error(f"Error uploading training data: {ex}")
        return templates.TemplateResponse("upload_training.html", {
            "request": request,
            "error": True,
            "message": f"Error processing file: {str(ex)}"
        })

@app.get("/view-training")
async def view_training_old(request: Request, db: Session = Depends(get_db)):
    """View training data"""
    try:
        from app.models.database import TrainingData
        training_data = db.query(TrainingData).order_by(TrainingData.created_at.desc()).limit(50).all()
        
        return templates.TemplateResponse("database_status.html", {
            "request": request,
            "training_data": training_data,
            "view_type": "training"
        })
    except Exception as ex:
        logger.error(f"Error retrieving training data: {ex}")
        return {"error": str(ex)}

@app.delete("/api/training/{training_id}")
async def delete_training(training_id: int, db: Session = Depends(get_db)):
    """Delete a training data entry"""
    try:
        from app.models.database import TrainingData
        
        # Find the training entry
        training_entry = db.query(TrainingData).filter(TrainingData.id == training_id).first()
        
        if not training_entry:
            return {"error": "Training data not found"}
        
        # Delete the entry
        db.delete(training_entry)
        db.commit()
        
        logger.info(f"Training data deleted: ID {training_id}")
        return {"message": "Training data deleted successfully"}
        
    except Exception as ex:
        logger.error(f"Error deleting training data: {ex}")
        db.rollback()
        return {"error": str(ex)}

@app.delete("/api/knowledge/{knowledge_id}")
async def delete_knowledge(knowledge_id: int, db: Session = Depends(get_db)):
    """Delete a knowledge base entry"""
    try:
        from app.models.database import KnowledgeBase
        
        # Find the knowledge entry
        knowledge_entry = db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_id).first()
        
        if not knowledge_entry:
            return {"error": "Knowledge entry not found"}
        
        # Delete the entry
        db.delete(knowledge_entry)
        db.commit()
        
        logger.info(f"Knowledge entry deleted: ID {knowledge_id}")
        return {"message": "Knowledge entry deleted successfully"}
        
    except Exception as ex:
        logger.error(f"Error deleting knowledge entry: {ex}")
        db.rollback()
        return {"error": str(ex)}

# Smart Solutions Routes

@app.get("/smart-solutions")
async def smart_solutions_get(request: Request):
    """Smart solutions page - GET"""
    return templates.TemplateResponse("smart_solutions.html", {"request": request})

@app.post("/smart-solutions")
async def smart_solutions_post(
    request: Request,
    problem_statement: str = Form(...),
    db: Session = Depends(get_db)
):
    """Find solutions based on error type matching - POST"""
    try:
        # Extract error type from problem statement
        error_matcher = ErrorTypeMatcher(db)
        error_type = await error_matcher.extract_and_store_error_type(problem_statement)
        
        # Find matching solutions
        solutions = await error_matcher.find_matching_solutions(error_type)
        
        logger.info(f"Found solutions for error type '{error_type}': {solutions['total_solutions']} total")
        
        return templates.TemplateResponse("smart_solutions.html", {
            "request": request,
            "problem_statement": problem_statement,
            "error_type": error_type,
            "solutions": solutions,
            "total_solutions": solutions['total_solutions']
        })
        
    except Exception as ex:
        logger.error(f"Error finding smart solutions: {ex}")
        return templates.TemplateResponse("smart_solutions.html", {
            "request": request,
            "problem_statement": problem_statement,
            "error": f"Error finding solutions: {str(ex)}"
        })

@app.post("/api/mark-useful")
async def mark_solution_useful(
    request: Request,
    db: Session = Depends(get_db)
):
    """Mark a solution as useful (API endpoint)"""
    try:
        data = await request.json()
        
        error_type = data.get("error_type")
        solution_type = data.get("solution_type")
        solution_id = data.get("solution_id")
        problem_statement = data.get("problem_statement")
        solution_title = data.get("solution_title", "")
        
        # Validate required fields
        if not all([error_type, solution_type, solution_id, problem_statement]):
            return {"error": "Missing required fields"}
        
        # Mark solution as useful
        error_matcher = ErrorTypeMatcher(db)
        success = await error_matcher.mark_solution_useful(
            error_type=error_type,
            solution_type=solution_type,
            solution_id=int(solution_id),
            problem_statement=problem_statement,
            user_id="web_user",
            feedback_notes=f"Marked useful for solution: {solution_title}"
        )
        
        if success:
            logger.info(f"Solution marked as useful: {solution_type}:{solution_id} for error_type '{error_type}'")
            return {"message": "Solution marked as useful", "success": True}
        else:
            return {"error": "Failed to mark solution as useful", "success": False}
        
    except Exception as ex:
        logger.error(f"Error marking solution as useful: {ex}")
        return {"error": f"Error: {str(ex)}", "success": False}

@app.get("/api/solution-details/{solution_type}/{solution_id}")
async def get_solution_details(
    solution_type: str,
    solution_id: int,
    db: Session = Depends(get_db)
):
    """Get full details of a solution (API endpoint)"""
    try:
        if solution_type == "knowledge_base":
            from app.models.database import KnowledgeBase
            solution = db.query(KnowledgeBase).filter(KnowledgeBase.id == solution_id).first()
            
            if not solution:
                return {"error": "Knowledge base solution not found"}
            
            return {
                "id": solution.id,
                "title": solution.title,
                "content": solution.content,
                "type": solution.type,
                "category": solution.category,
                "keywords": solution.keywords,
                "priority": solution.priority,
                "created_at": solution.created_at.isoformat() if solution.created_at else None,
                "solution_type": "knowledge_base"
            }
            
        elif solution_type == "incident_case":
            from app.models.database import TrainingData
            solution = db.query(TrainingData).filter(TrainingData.id == solution_id).first()
            
            if not solution:
                return {"error": "Incident case solution not found"}
            
            return {
                "id": solution.id,
                "title": f"Incident Case: {solution.expected_incident_type}",
                "content": solution.incident_description,
                "resolution": solution.expected_root_cause,
                "impact": solution.expected_impact,
                "urgency": solution.expected_urgency,
                "affected_systems": solution.expected_affected_systems,
                "category": solution.category,
                "created_at": solution.created_at.isoformat() if solution.created_at else None,
                "solution_type": "incident_case"
            }
        else:
            return {"error": "Invalid solution type"}
        
    except Exception as ex:
        logger.error(f"Error getting solution details: {ex}")
        return {"error": f"Error: {str(ex)}"}

@app.get("/api/error-analytics")
async def get_error_analytics(
    error_type: str = None,
    db: Session = Depends(get_db)
):
    """Get analytics for error types and solution effectiveness"""
    try:
        error_matcher = ErrorTypeMatcher(db)
        analytics = await error_matcher.get_error_type_analytics(error_type)
        return analytics
        
    except Exception as ex:
        logger.error(f"Error getting analytics: {ex}")
        return {"error": f"Error: {str(ex)}"}

# Document Parsing API Endpoints
@app.post("/api/parse-bulk-content")
async def parse_bulk_content(
    request: Request,
    bulk_content: str = Form(...),
    content_type: str = Form("knowledge_base"),
    default_category: str = Form("General"),
    db: Session = Depends(get_db)
):
    """Parse bulk pasted content into individual entries using AI"""
    try:
        logger.info(f"Parsing bulk content of {len(bulk_content)} characters")
        
        # Parse content using AI
        entries = await document_parser.parse_bulk_content(bulk_content, content_type)
        
        # Apply default category if not set by AI
        for entry in entries:
            if not entry.get('category') or entry['category'] == 'General':
                entry['category'] = default_category
        
        logger.info(f"Successfully parsed {len(entries)} entries")
        
        return {
            "status": "success",
            "entries": entries,
            "message": f"Successfully parsed {len(entries)} entries from content"
        }
        
    except Exception as ex:
        logger.error(f"Error parsing bulk content: {ex}")
        return {"status": "error", "error": f"Error parsing content: {str(ex)}"}

@app.post("/api/parse-uploaded-file")
async def parse_uploaded_file(
    request: Request,
    document_file: UploadFile = File(...),
    upload_content_type: str = Form("knowledge_base"),
    upload_category: str = Form("General"),
    db: Session = Depends(get_db)
):
    """Parse uploaded document file into individual entries"""
    try:
        # Validate file size (10MB limit)
        if document_file.size > 10 * 1024 * 1024:
            return {"status": "error", "error": "File size exceeds 10MB limit"}
        
                        # Validate file type
        allowed_extensions = {'.docx', '.pdf', '.txt'}  # Removed .doc for now
        file_extension = os.path.splitext(document_file.filename)[1].lower()
        if file_extension not in allowed_extensions:
            return {"status": "error", "error": f"Unsupported file type: {file_extension}. Supported formats: .docx, .pdf, .txt"}
        
        logger.info(f"Processing uploaded file: {document_file.filename}")
        
        # Read file content
        file_content = await document_file.read()
        
        # Parse file using AI
        entries = await document_parser.parse_file_content(file_content, document_file.filename)
        
        # Apply default category
        for entry in entries:
            if not entry.get('category') or entry['category'] == 'General':
                entry['category'] = upload_category
        
        logger.info(f"Successfully parsed {len(entries)} entries from file")
        
        return {
            "status": "success",
            "entries": entries,
            "message": f"Successfully parsed {len(entries)} entries from {document_file.filename}"
        }
        
    except Exception as ex:
        logger.error(f"Error parsing uploaded file: {ex}")
        return {"status": "error", "error": f"Error parsing file: {str(ex)}"}

@app.post("/api/save-bulk-entries")
async def save_bulk_entries(
    request: Request,
    db: Session = Depends(get_db)
):
    """Save multiple parsed entries to the knowledge base"""
    try:
        data = await request.json()
        entries = data.get("entries", [])
        
        if not entries:
            return {"status": "error", "error": "No entries provided"}
        
        from app.models.database import KnowledgeBase
        
        saved_count = 0
        for entry_data in entries:
            try:
                # Create new knowledge base entry
                kb_entry = KnowledgeBase(
                    title=entry_data.get('title', '').strip()[:200],
                    content=entry_data.get('content', entry_data.get('solution', entry_data.get('description', ''))).strip()[:5000],
                    category=entry_data.get('category', 'General').strip()[:100],
                    type='Solution',
                    tags=entry_data.get('tags', '').strip()[:500],
                    keywords=entry_data.get('keywords', '').strip()[:500],
                    priority=entry_data.get('priority', 'Medium'),
                    source='Document Import',
                    status='Active',
                    view_count=0,
                    created_at=datetime.now(),
                    updated_at=datetime.now(),
                    created_by='AI Assistant'
                )
                
                db.add(kb_entry)
                saved_count += 1
                
            except Exception as entry_error:
                logger.warning(f"Failed to save entry: {entry_error}")
                continue
        
        db.commit()
        
        logger.info(f"Successfully saved {saved_count} entries to knowledge base")
        
        return {
            "status": "success",
            "saved_count": saved_count,
            "message": f"Successfully saved {saved_count} entries to the knowledge base"
        }
        
    except Exception as ex:
        logger.error(f"Error saving bulk entries: {ex}")
        db.rollback()
        return {"status": "error", "error": f"Error saving entries: {str(ex)}"}

# Knowledge Base Management Route
@app.get("/knowledge-base-manage")
async def knowledge_base_manage(request: Request):
    """Enhanced knowledge base management page with AI parsing"""
    return templates.TemplateResponse("knowledge_base_manage.html", {
        "request": request,
        "title": "Knowledge Base Management"
    })

@app.post("/api/debug-document-content")
async def debug_document_content(
    request: Request,
    document_file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """Debug endpoint to show raw extracted content from document with AI analysis"""
    try:
        # Read file content
        file_content = await document_file.read()
        
        # Extract text content using document parser
        extracted_text = ""
        content_parts = []
        
        if document_file.filename.lower().endswith('.docx'):
            from docx import Document
            import io
            
            doc = Document(io.BytesIO(file_content))
            
            # Extract all text with structure info
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    style = paragraph.style.name if paragraph.style else "Normal"
                    content_parts.append({
                        "type": "paragraph",
                        "style": style,
                        "text": paragraph.text.strip()
                    })
                    extracted_text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text = " | ".join(row_text)
                        content_parts.append({
                            "type": "table_row",
                            "text": table_text
                        })
                        extracted_text += table_text + "\n"
        
        elif document_file.filename.lower().endswith(('.txt', '.pdf')):
            # Use document parser service for other formats
            extracted_text = await document_parser._extract_from_pdf(file_content) if document_file.filename.lower().endswith('.pdf') else file_content.decode('utf-8')
            content_parts = [{"type": "text", "text": extracted_text}]
        
        else:
            return {"status": "error", "error": "Supported formats: .docx, .pdf, .txt"}
        
        # Use AI analysis from document parser service
        structure_analysis = document_parser.analyze_document_structure(extracted_text)
        
        return {
            "status": "success",
            "filename": document_file.filename,
            "file_size_bytes": len(file_content),
            
            # Structure from document format
            "document_structure": {
                "total_paragraphs": len([p for p in content_parts if p["type"] == "paragraph"]),
                "total_tables": len([p for p in content_parts if p["type"] == "table_row"]),
                "content_structure": content_parts[:20],  # First 20 items
            },
            
            # AI analysis from document parser
            "ai_analysis": structure_analysis,
            
            # Raw text info
            "raw_text_info": {
                "total_length": len(extracted_text),
                "preview": extracted_text[:1500] + "..." if len(extracted_text) > 1500 else extracted_text,
                "lines_count": len(extracted_text.splitlines()),
            },
            
            # Extraction potential estimate
            "extraction_estimate": {
                "expected_entries": max(1, structure_analysis.get("potential_sections", 0)),
                "confidence": "high" if structure_analysis.get("potential_sections", 0) > 5 else "medium" if structure_analysis.get("potential_sections", 0) > 2 else "low"
            }
        }
            
    except Exception as ex:
        logger.error(f"Error debugging document: {ex}")
        return {"status": "error", "error": f"Error: {str(ex)}"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="localhost", port=8001)
